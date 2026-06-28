#!/usr/bin/env python3
"""Compare a filled-in 様式 docx against its blank template; flag fill mistakes.

Usage:
    diff-form-docx.py <filled.docx> <blank.docx> [--strict]
    diff-form-docx.py --selftest

The docx analog of `diff-form-xlsx.py`. Word 様式 (cover sheets, 申請書) have no
mechanical fill-check, so a whole class of mistakes slips through even after a
human renders the PDF and "looks at it" — because the human looks at *what they
filled*, not at *what a reviewer would see as blank / overwritten / mis-placed*.

Findings
--------
HARD (exit 1):
  ❌ LABEL_OVERWRITE  : a table cell that held pre-printed label text in the blank
                        form now holds *different* text (no containment either way).
                        = the cover-sheet bug: writing a value into a LABEL cell
                          (e.g. birth date typed into the 「年令」/age label cell).
  ❌ LABEL_DELETED    : a blank label cell is empty in the filled form.
  ❌ HEADING_VANISHED : a structural heading paragraph ( （N） / N． ) present in the
                        blank form is absent from the filled form (= overwritten /
                        deleted, e.g. a heading clobbered by an off-by-one fill).

SURFACE (printed; exit 0 unless --strict):
  ⚠ BULLET_EMPTY     : a bullet placeholder ( 「・」 alone ) is still empty
                        = a whole 「・・・」 section left unfilled.
  ⚠ EMPTY_LABELED_COL: a table column with a non-empty header has *all* data cells
                        empty = a labeled column (e.g. 「年月」) left blank.
  ℹ LABEL_PARTIAL    : filled text is a substring of the blank label (label split /
                        restructured, e.g. one 「学歴・職歴」 cell broken into rows).
  ℹ INPUT_FILLED     : blank cell, filled has text (= normal applicant input).

The check is structural, not semantic: it can tell that a box / column / bullet is
EMPTY or that a label was REPLACED, but it cannot judge whether prose content is
*thin* — that remains a human "reviewer's-eye" pass (see
office-automation-principles.md#reviewer-eye-completion).

Why this script exists
----------------------
2026-06-27 ある研究費 docx 申請様式を 1 session で記入した際、同じ様式に対し
4 つの記入ミスを連続で出し、そのつど人間 (申請者) が見つけて指摘した:
  (1) 生年月日を「年令」ラベルセルに上書き  → LABEL_OVERWRITE
  (2) 「（１）研究者氏名・所属先」の「・・・」を空欄のまま  → BULLET_EMPTY
  (3) 略歴の「年月」列を空のまま (内容列に畳んだ)  → EMPTY_LABELED_COL
  (4) ≪研究内容概要≫ の箱が下半分・薄い  → (箱の空きは検出可、薄さは判断)
さらに自分で catch した off-by-one 見出し上書き (= 「（５）...」見出しの clobber) も
HEADING_VANISHED として機械化する。xlsx には diff-form-xlsx.py があるが docx には
無かった = この session の記入が機械の網ゼロだった。これを埋める。
"""

import re
import sys
import unicodedata

try:
    from docx import Document
except ImportError:  # pragma: no cover
    print("python-docx required: pip install python-docx", file=sys.stderr)
    sys.exit(2)

# Heading-like paragraph: （１） / (1) / １． / 1. / １、 etc. — pure structural labels
# that must survive verbatim in the filled form.
HEADING_RE = re.compile(r"^[\s　]*(?:[（(][0-9０-９]+[）)]|[0-9０-９]+[．.、])")
# A paragraph whose normalized text is exactly one of these is an empty bullet slot.
BULLET_TOKENS = {"・", "·", "•", "‧", "・", "-", "*", "／", "/"}


def norm(s):
    """NFKC + strip ALL whitespace (incl. full-width 　 and newlines)."""
    if s is None:
        return ""
    s = unicodedata.normalize("NFKC", str(s))
    return re.sub(r"\s+", "", s)


def all_text_blocks(doc):
    """Every text-bearing block (body paragraphs + all table cells), flattened."""
    blocks = []
    for p in doc.paragraphs:
        blocks.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                blocks.append(cell.text)
    return blocks


def diff_forms(filled_path, blank_path):
    fdoc = Document(filled_path)
    bdoc = Document(blank_path)
    findings = {k: [] for k in (
        "LABEL_OVERWRITE", "LABEL_DELETED", "HEADING_VANISHED",
        "BULLET_EMPTY", "EMPTY_LABELED_COL", "LABEL_PARTIAL", "INPUT_FILLED")}

    _diff_tables(fdoc, bdoc, findings)
    _check_headings(fdoc, bdoc, findings)
    _check_bullets(fdoc, findings)
    return findings


def _diff_tables(fdoc, bdoc, findings):
    for ti, (bt, ft) in enumerate(zip(bdoc.tables, fdoc.tables)):
        nrow = min(len(bt.rows), len(ft.rows))
        ncol = min(len(bt.columns), len(ft.columns))
        # --- per-cell label / input diff ---
        for r in range(nrow):
            for c in range(ncol):
                try:
                    bv = bt.cell(r, c).text
                    fv = ft.cell(r, c).text
                except (IndexError, ValueError):
                    continue
                bn, fn = norm(bv), norm(fv)
                if bn == fn:
                    continue
                loc = f"table{ti}!r{r}c{c}"
                if not bn:                       # blank empty -> applicant input
                    if fn:
                        findings["INPUT_FILLED"].append((loc, _short(fv)))
                elif not fn:                     # label cell now empty
                    findings["LABEL_DELETED"].append((loc, f"blank={_short(bv)} -> empty"))
                elif bn in fn:                   # label preserved + appended (e.g. ○第１課題)
                    pass                         # OK (mark / box label kept)
                elif fn in bn:                   # filled is a subset of the label (restructure)
                    findings["LABEL_PARTIAL"].append((loc, f"blank={_short(bv)} -> {_short(fv)}"))
                else:                            # text fully replaced = overwrite
                    findings["LABEL_OVERWRITE"].append((loc, f"blank={_short(bv)} -> {_short(fv)}"))
        # --- labeled column with all data cells empty (in the FILLED form) ---
        # Only for "header tables" = blank row 0 is a *full* header (every cell
        # non-empty). This excludes key-value cover sheets (label|value|label|…
        # where blank row 0 has gaps at the value cells), so 略歴/予算表 (uniform
        # header row) are checked but the 表紙 is not.
        blank_header_full = (
            len(bt.columns) >= 2
            and all(norm(bt.cell(0, c).text) for c in range(len(bt.columns)))
        )
        if blank_header_full and len(ft.rows) >= 2 and len(ft.columns) >= 1:
            for c in range(len(ft.columns)):
                try:
                    header = norm(ft.cell(0, c).text)
                except (IndexError, ValueError):
                    continue
                if not header:
                    continue
                data_cells = []
                seen = set()
                for r in range(1, len(ft.rows)):
                    try:
                        cell = ft.cell(r, c)
                    except (IndexError, ValueError):
                        continue
                    key = id(cell._tc)           # dedupe merged cells
                    if key in seen:
                        continue
                    seen.add(key)
                    data_cells.append(norm(cell.text))
                if data_cells and all(not d for d in data_cells):
                    findings["EMPTY_LABELED_COL"].append(
                        (f"table{ti} col{c}", f"header={_short(ft.cell(0, c).text)} / data cells all empty"))


def _check_headings(fdoc, bdoc, findings):
    filled_text = norm("".join(all_text_blocks(fdoc)))
    for p in bdoc.paragraphs:
        raw = p.text
        if not HEADING_RE.match(raw):
            continue
        hn = norm(raw)
        if not hn:
            continue
        if hn not in filled_text:
            findings["HEADING_VANISHED"].append(("body", _short(raw)))


def _check_bullets(fdoc, findings):
    for p in fdoc.paragraphs:
        if norm(p.text) in BULLET_TOKENS:
            findings["BULLET_EMPTY"].append(("body", f"empty bullet: {p.text!r}"))
    for t in fdoc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if norm(p.text) in BULLET_TOKENS:
                        findings["BULLET_EMPTY"].append(("table-cell", f"empty bullet: {p.text!r}"))


def _short(v, n=50):
    s = str(v).replace("\n", "⏎")
    return f"{s[:n]}{'…' if len(s) > n else ''}"


HARD = ("LABEL_OVERWRITE", "LABEL_DELETED", "HEADING_VANISHED")
SOFT = ("BULLET_EMPTY", "EMPTY_LABELED_COL", "LABEL_PARTIAL")


def report(findings, strict=False):
    hard = soft = 0
    for cls in HARD + SOFT + ("INPUT_FILLED",):
        items = findings[cls]
        if not items:
            continue
        if cls in HARD:
            marker = "❌"
        elif cls == "INPUT_FILLED":
            marker = "ℹ"
        else:
            marker = "⚠"
        print(f"\n{marker} {cls} ({len(items)})")
        for loc, detail in items[:40]:
            print(f"  {loc}: {detail}")
        if len(items) > 40:
            print(f"  … +{len(items) - 40} more")
        if cls in HARD:
            hard += len(items)
        elif cls in SOFT:
            soft += len(items)

    if hard:
        print(f"\n❌ {hard} HARD finding(s): 様式の label/見出しが上書き・消失している。")
        print("   ラベルセルは入力欄ではない — 元のラベルを復元し、値は値セル側へ。")
    if soft:
        print(f"\n⚠ {soft} 完成度 finding(s): 空の箇条書き / 全空の列 / 部分改変。")
        print("   意図的空欄 (印・提出日・指導者欄 等) か、記入漏れかを 1 件ずつ確認。")
    if not hard and not soft:
        print("\n✓ 上書き・見出し消失・空の箇条書き/列 は検出されませんでした。")
        print("  ⚠ ただし「中身が薄い/箱が半分」 は機械では判定不可 — 審査員の目で最終確認。")

    if hard or (strict and soft):
        return 1
    return 0


# --------------------------------------------------------------------------- #
# selftest: synthetic blank/filled docx reproducing each 2026-06-27 mistake
# --------------------------------------------------------------------------- #
def _selftest():
    import tempfile, os
    from docx import Document as D

    def build(path, kind):
        d = D()
        # cover-like table: r0 ふりがな, r1 氏名|印|年令|歳, ...
        t = d.add_table(rows=2, cols=4)
        t.cell(0, 0).text = "申請者氏名"
        t.cell(0, 2).text = "年\n令"            # age LABEL (the trap)
        t.cell(0, 3).text = "歳"
        t.cell(1, 0).text = "応募課題"
        t.cell(1, 1).text = "第１課題　第２課題　（いずれかを○で）"
        if kind == "blank":
            pass                                 # value cells empty
        elif kind == "good":
            t.cell(0, 1).text = "甲野 太郎"      # name -> value cell (c1) OK
            t.cell(1, 1).text = "○第１課題　第２課題　（いずれかを○で）"  # mark prepended OK
        elif kind == "overwrite":
            t.cell(0, 2).text = "1900.1.1生"     # DOB typed into 年令 LABEL = bug (1)
        # 略歴-like table: header 項目/年月/内容
        t2 = d.add_table(rows=2, cols=3)
        t2.cell(0, 0).text = "項目"
        t2.cell(0, 1).text = "年月"
        t2.cell(0, 2).text = "内容"
        if kind == "good":
            t2.cell(1, 0).text = "学歴"
            t2.cell(1, 1).text = "1990. 3"       # 年月 column filled
            t2.cell(1, 2).text = "甲野大学 卒業"
        elif kind == "emptycol":
            t2.cell(1, 0).text = "学歴"
            t2.cell(1, 2).text = "甲野大学 卒業"  # 年月 (c1) left empty = bug (3)
        # heading + bullet section
        d.add_paragraph("（１）研究者氏名・所属先")
        if kind in ("blank", "emptycol", "overwrite"):
            d.add_paragraph("・")                # empty bullet = bug (2)
            d.add_paragraph("・")
        elif kind == "good":
            d.add_paragraph("・甲野 太郎　甲野大学")
        d.add_paragraph("（５）研究の開始・終了予定年月日")  # heading that must survive
        if kind == "vanished":
            # rebuild WITHOUT the (5) heading to simulate clobber
            pass
        d.save(path)

    tmp = tempfile.mkdtemp()
    paths = {}
    for k in ("blank", "good", "overwrite", "emptycol"):
        p = os.path.join(tmp, f"{k}.docx")
        build(p, k)
        paths[k] = p
    # vanished: filled good but with (5) heading removed
    vdoc = Document(paths["good"])
    for para in list(vdoc.paragraphs):
        if para.text.startswith("（５）"):
            para._element.getparent().remove(para._element)
    vpath = os.path.join(tmp, "vanished.docx")
    vdoc.save(vpath)
    paths["vanished"] = vpath

    def has(f, cls):
        return len(diff_forms(paths[f], paths["blank"])[cls]) > 0

    checks = [
        ("good: no LABEL_OVERWRITE", not has("good", "LABEL_OVERWRITE")),
        ("good: no BULLET_EMPTY", not has("good", "BULLET_EMPTY")),
        ("good: no EMPTY_LABELED_COL", not has("good", "EMPTY_LABELED_COL")),
        ("good: no HEADING_VANISHED", not has("good", "HEADING_VANISHED")),
        ("overwrite: LABEL_OVERWRITE detected (年令)", has("overwrite", "LABEL_OVERWRITE")),
        ("blank: BULLET_EMPTY detected", has("blank", "BULLET_EMPTY")),
        ("emptycol: EMPTY_LABELED_COL detected (年月)", has("emptycol", "EMPTY_LABELED_COL")),
        ("emptycol: BULLET_EMPTY detected", has("emptycol", "BULLET_EMPTY")),
        ("vanished: HEADING_VANISHED detected ((5))", has("vanished", "HEADING_VANISHED")),
        ("good: INPUT_FILLED present (name)", has("good", "INPUT_FILLED")),
    ]
    ok = True
    for name, passed in checks:
        print(f"  {'PASS' if passed else 'FAIL'}  {name}")
        ok = ok and passed
    print(f"\n{'ALL PASS' if ok else 'FAILURES'} ({sum(p for _, p in checks)}/{len(checks)})")
    return 0 if ok else 1


def main():
    args = sys.argv[1:]
    if "--selftest" in args:
        sys.exit(_selftest())
    strict = "--strict" in args
    args = [a for a in args if not a.startswith("--")]
    if len(args) != 2:
        print(__doc__, file=sys.stderr)
        sys.exit(2)
    findings = diff_forms(args[0], args[1])
    sys.exit(report(findings, strict=strict))


if __name__ == "__main__":
    main()
